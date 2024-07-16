import base64
import os
import shutil
import time
from docx import Document
from io import BytesIO
from openai import OpenAI
import streamlit as st
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.core.os_manager import ChromeType
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from PIL import Image
import io


def capture_full_page_screenshots(url, output_folder):
    # Delete the output folder if it already exists
    if os.path.exists(output_folder):
        shutil.rmtree(output_folder)
    
    # Create the output folder
    os.makedirs(output_folder)

    options = Options()
    options.add_argument("--disable-gpu")
    options.add_argument("--headless")
    
    # Set up the webdriver (Chromium in this example)
    driver = webdriver.Chrome(
            service=Service(
                ChromeDriverManager(chrome_type=ChromeType.CHROMIUM).install()
            ),
            options=options,
        )
    
    try:
        # Open the URL
        driver.get(url)
        
        # Wait for the page to load
        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        
        # Get the page dimensions
        total_height = driver.execute_script("return document.body.scrollHeight")
        viewport_height = driver.execute_script("return window.innerHeight")
        
        # Calculate the number of screenshots needed
        num_screenshots = total_height // viewport_height + 1
        
        screenshot_paths = []
        
        for i in range(num_screenshots):
            # Calculate scroll position
            scroll_height = i * viewport_height
            if i > 0:
                scroll_height -= 100  # Overlap by 100 pixels
            
            # Scroll to position
            driver.execute_script(f"window.scrollTo(0, {scroll_height});")
            
            # Wait for any dynamic content to load
            time.sleep(1)
            
            # Capture the screenshot
            screenshot = driver.get_screenshot_as_base64()
            img_data = base64.b64decode(screenshot)
            img = Image.open(io.BytesIO(img_data))
            
            # Save the screenshot
            filename = f"screenshot_{i+1:03d}.png"
            file_path = os.path.join(output_folder, filename)
            img.save(file_path)
            screenshot_paths.append(file_path)
            
            print(f"Captured and saved screenshot {i+1} of {num_screenshots}")
        
        print(f"All screenshots saved in folder: {output_folder}")
        
        return screenshot_paths
        
    finally:
        # Close the browser
        driver.quit()


def analyze_all_screenshots(api_key,screenshot_paths):
    client = OpenAI(api_key=api_key)
    
    base64_images = []
    for path in screenshot_paths:
        with open(path, "rb") as image_file:
            base64_images.append(base64.b64encode(image_file.read()).decode('utf-8'))
    
    messages = [
        
    {
        "role": "system",
        "content": """You are an expert in website user interface (UI) and user experience (UX) design, as well as content analysis. Your task is to analyze screenshots of websites and provide detailed, structured reports on their UI/UX and content quality. Your analysis should be thorough and based on the following criteria:

        1. **UI/UX Evaluation**:
            - **Navigation Structure and Usability**: Assess how easy it is to navigate the website, including menu design and link placement.
            - **Layout and Visual Hierarchy**: Evaluate the arrangement of elements on the page and their visual importance.
            - **Accessibility Features**: Identify any accessibility features present or missing, such as alt text for images, color contrast, and keyboard navigation.
            - **Visual Design and Appeal**: Comment on the aesthetic aspects of the UI, including color schemes, typography, iconography, and overall design cohesion.
            - **Interaction Design**: Analyze the responsiveness and feedback of interactive elements like buttons, forms, and hover effects.
            - **Call-to-Action Effectiveness**: Examine the clarity, visibility, and placement of calls-to-action (CTAs).
            - **Consistency**: Evaluate the consistency of UI elements across different sections of the website.

        2. **Content Quality Assessment**:
            - **Content Clarity and Readability**: Evaluate how clear and easy to read the content is, considering language, font size, and formatting.
            - **Relevance to Target Audience**: Assess whether the content is appropriate and engaging for the intended audience.
            - **Engagement Factors**: Review the tone and style of the content, as well as the use of multimedia (images, videos, etc.) to engage users.
            - **Call-to-Action Effectiveness**: Check the effectiveness of CTAs in the content, their wording, and positioning.
            - **Branding and Messaging Consistency**: Look for consistency in branding elements such as logos, color schemes, and messaging.
            - **Content Organization and Structure**: Assess the logical flow and organization of the content, including headings, subheadings, and paragraphs.

        For each criterion, clearly state any issues found, their potential impact on user interface, user experience, and content effectiveness, and provide actionable recommendations for improvement. The final report should include a summary of key findings, prioritization of issues based on impact and effort required to fix them, and suggested key performance indicators (KPIs) to track improvements."""
    },
        
        {
            "role": "user",
            "content": [
                {"type": "text", "text": """Analyze these screenshots of a website for UI/UX and content.
                 
                Evaluate the user interface and user experience of the website. Provide a structured report with the following sections:
                1. Navigation Structure and Usability
                2. Layout and Visual Hierarchy
                3. Accessibility Features
                4. Visual Design and Appeal
                5. Interaction Design
                6. Call-to-Action Effectiveness
                7. Consistency
                For each section, clearly state the issues found and their potential impact on user interface and user experience.
                Expected Output=A structured UI/UX evaluation report with clear sections, 
                each detailing specific issues found and their potential impact on user interface and user experience.
                
                Assess the content quality and effectiveness of the website.
                Provide a structured report with the following sections:
                1. Content Clarity and Readability
                2. Relevance to Target Audience
                3. Engagement Factors (tone, style, multimedia use)
                4. Call-to-Action Effectiveness
                5. Branding and Messaging Consistency
                6. Content Organization and Structure
                For each section, clearly state the issues found and their potential impact on content effectiveness.
                Expected Output=A structured content assessment report with clear sections, 
                each detailing specific issues found and their potential impact on content effectiveness.
                
                Review the UI/UX evaluation and content assessment.
                Create a comprehensive report that addresses all identified issues. Your report should:
                1. Summarize key findings from each analysis (UI/UX and Content)
                2. For each issue identified:
                    a. Clearly state the problem
                    b. Explain its impact on the website's performance
                    c. Provide a specific, actionable solution
                    d. Estimate the effort required (Low/Medium/High)
                    e. Predict the potential impact of implementing the solution (Low/Medium/High)
                3. Prioritize the solutions based on their potential impact and required effort
                4. Propose an implementation timeline
                5. Suggest key performance indicators to track improvement
                expected_output=A structured content assessment report with clear sections, 
                each detailing specific issues found and their potential impact on content effectiveness.
                
                Also tell me total number of images analyzed by you"""},
            ] + [{"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img}"}} for img in base64_images]
        }
    ]
    
    response = client.chat.completions.create(
        model="gpt-4o", # Make sure to use a vision-capable model
        messages=messages,
        max_tokens=4000,
    )
    
    return response.choices[0].message.content

# Streamlit web application
def main():
    st.header('AI Website UX and Content Critic')
    
    # Initialize session state
    if 'generated_content' not in st.session_state:
        st.session_state.generated_content = None

    with st.sidebar:
        with st.form('OpenAI'):
            api_key = st.text_input('Enter your OpenAI API key', type="password")
            submitted = st.form_submit_button("Submit")


    if api_key:
        
        url = st.text_input("Enter your URL")
        
        if st.button("Critique It!"):
            with st.spinner("Capturing screenshots..."):
                output_folder = "webpage_screenshots"
                screenshot_paths = capture_full_page_screenshots(url, output_folder)
            
            with st.spinner("Critiquing..."):
                st.session_state.generated_content = analyze_all_screenshots(api_key,screenshot_paths)

        # Display content if it exists in session state
        if st.session_state.generated_content:
            st.markdown(st.session_state.generated_content)

            doc = Document()

            # Option to download content as a Word document
            doc.add_heading(url, 0)
            doc.add_paragraph(st.session_state.generated_content)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="Download as Word Document",
                data=buffer,
                file_name=f"{url}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
